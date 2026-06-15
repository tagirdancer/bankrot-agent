"""
Обнаружение, скачивание и разбор всех документов лота (PDF, DOCX, фото).
"""
from __future__ import annotations

import io
import logging
import re
from typing import Any
from urllib.parse import urljoin, urlparse

log = logging.getLogger("lot_documents")

IMAGE_EXTS = frozenset({"jpg", "jpeg", "png", "gif", "webp", "bmp"})
READABLE_EXTS = frozenset({"pdf", "docx", "doc"})


def _norm_url(url: str, base: str = "https://tbankrot.ru") -> str:
    url = (url or "").strip().replace("\\/", "/")
    if url.startswith("//"):
        return "https:" + url
    if url.startswith("http"):
        return url.split("#")[0]
    return urljoin(base + "/", url.lstrip("/"))


def _file_ext(url: str, title: str = "") -> str:
    path = urlparse(url).path.lower()
    for src in (path, title.lower()):
        m = re.search(r"\.(pdf|docx?|jpe?g|png|gif|webp|bmp)(?:\?|$)", src)
        if m:
            return m.group(1).replace("jpeg", "jpg")
    if "file-store" in url and ".docx" in title.lower():
        return "docx"
    if "file-store" in url and ".pdf" in title.lower():
        return "pdf"
    if "etpphoto" in url.lower() or "photo" in url.lower():
        return "jpg"
    return ""


def discover_lot_documents(html: str, lot_id: str) -> list[dict]:
    """
    Все документы из раздела «Документы по торгам» и смежных блоков карточки.
    Возвращает [{url, title, ext}, ...] без дубликатов.
    """
    if not html:
        return []

    docs: list[dict] = []
    seen_urls: set[str] = set()

    def add(url: str, title: str) -> None:
        url = _norm_url(url)
        if not url or url in seen_urls:
            return
        if not _is_document_url(url, title):
            return
        seen_urls.add(url)
        docs.append({
            "url": url,
            "title": re.sub(r"\s+", " ", (title or "").strip())[:240],
            "ext": _file_ext(url, title),
        })

    # Якоря в таблице документов и рядом
    for m in re.finditer(
        r'<a[^>]+href=["\']([^"\']+)["\'][^>]*>([^<]{2,240})</a>',
        html, re.I,
    ):
        add(m.group(1), m.group(2))

    # file-store / egrn / фото без текста ссылки
    for m in re.finditer(
        r'https?://(?:torgi\.gov\.ru/new/file-store/v1/[a-f0-9]+|'
        r'files\.tbankrot\.ru/(?:egrn_files|etpPhoto)/[^"\'>\s]+)',
        html, re.I,
    ):
        add(m.group(0), "")

    # get_doc.php шаблоны (если есть в блоке документов)
    for m in re.finditer(
        r'get_doc\.php\?[^"\']+',
        html, re.I,
    ):
        add(_norm_url(m.group(0)), f"шаблон {lot_id}")

    return docs


def _is_document_url(url: str, title: str) -> bool:
    ul = url.lower()
    tl = (title or "").lower()
    if any(x in ul for x in (
        "file-store", "egrn_files", "etpphoto", "get_doc.php",
        ".pdf", ".docx", ".doc", ".jpg", ".jpeg", ".png",
    )):
        return True
    if any(x in tl for x in (".pdf", ".docx", ".doc", ".jpg", ".jpeg", ".png")):
        return True
    if re.search(r"егрн|оценк|договор|заявк|инф\.?\s*сообщ", tl):
        return True
    return False


def classify_document(title: str, url: str, text: str, ext: str) -> str:
    """egrn | appraisal | contract | application | info_message | photo | other"""
    if ext in IMAGE_EXTS:
        return "photo"

    blob = f"{title} {url}".lower()
    tl = (text or "")[:4000].lower()

    if ext in ("docx", "doc"):
        if re.search(r"заявк", blob):
            return "application"
        if re.search(r"договор", blob):
            return "contract"
        return "other"

    if re.search(r"инф\.?\s*сообщ|информационн[^\n]{0,20}сообщ", blob):
        return "info_message"
    if re.search(r"егрн|выписк[^\n]{0,40}егрн", blob):
        return "egrn"
    if re.search(r"оцен|отч[её]т[^\n]{0,30}оцен", blob):
        return "appraisal"

    if text:
        try:
            from appraisal_pdf import classify_pdf_type
            ct = classify_pdf_type(text, url)
            if ct in ("egrn", "appraisal"):
                return ct
        except ImportError:
            pass
        if re.search(r"инф\.?\s*сообщ|информационн[^\n]{0,20}сообщ", tl):
            return "info_message"
        if re.search(r"договор[^\n]{0,40}купл", tl):
            return "contract"
        if re.search(r"заявк[^\n]{0,40}(?:участ|торг)", tl):
            return "application"

    return "other"


def extract_docx_text(raw: bytes) -> tuple[str, str]:
    try:
        from docx import Document
        doc = Document(io.BytesIO(raw))
        parts = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        text = "\n".join(parts).strip()
        if len(text) >= 40:
            return text[:20000], "docx"
    except Exception as e:
        log.warning("docx extract failed: %s", e)
    return "", "failed"


def parse_contract_text(text: str) -> dict[str, Any]:
    result = {
        "doc_type": "contract",
        "price": 0,
        "parties": "не указано",
        "object_description": "не указано",
        "parsed_ok": False,
        "summary": "",
    }
    if not text or len(text) < 40:
        return result
    norm = re.sub(r"[ \t]+", " ", text)

    pm = re.search(
        r"(\d[\d\s]{5,11})(?:[.,](\d{2}))?\s*(?:₽|руб\.?)",
        norm, re.I,
    )
    if pm:
        result["price"] = int(re.sub(r"\s", "", pm.group(1)))

    for pat in (
        r"продавец[:\s]+([^\n|]{5,120})",
        r"покупатель[:\s]+([^\n|]{5,120})",
    ):
        m = re.search(pat, norm, re.I)
        if m:
            result["parties"] = m.group(1).strip()[:160]
            break

    om = re.search(
        r"(?:предмет договора|объект)[:\s]+([^\n]{15,220})",
        norm, re.I,
    )
    if om:
        result["object_description"] = om.group(1).strip()[:200]

    result["parsed_ok"] = bool(result["price"] or result["parties"] != "не указано"
                               or result["object_description"] != "не указано")
    parts = []
    if result["price"]:
        parts.append(f"цена {result['price']/1e6:.2f} млн")
    if result["parties"] != "не указано":
        parts.append(f"стороны: {result['parties'][:80]}")
    if result["object_description"] != "не указано":
        parts.append(result["object_description"][:80])
    result["summary"] = "; ".join(parts) or "текст прочитан, ключевые поля не найдены"
    return result


def parse_application_text(text: str) -> dict[str, Any]:
    result = {
        "doc_type": "application",
        "applicant": "не указано",
        "deposit": "не указано",
        "parsed_ok": False,
        "summary": "",
    }
    if not text or len(text) < 40:
        return result
    norm = re.sub(r"[ \t]+", " ", text)

    am = re.search(
        r"(?:заявитель|участник торгов|ф\.?\s*и\.?\s*о\.?)[:\s]+([^\n|]{5,120})",
        norm, re.I,
    )
    if am:
        result["applicant"] = am.group(1).strip()[:160]

    dm = re.search(
        r"задат(?:ок|ка)[^\d]{0,30}(\d[\d\s]{3,11})\s*(?:₽|руб)",
        norm, re.I,
    )
    if dm:
        result["deposit"] = re.sub(r"\s", "", dm.group(1)) + " ₽"

    result["parsed_ok"] = result["applicant"] != "не указано" or result["deposit"] != "не указано"
    parts = []
    if result["applicant"] != "не указано":
        parts.append(f"заявитель: {result['applicant'][:80]}")
    if result["deposit"] != "не указано":
        parts.append(f"задаток: {result['deposit']}")
    result["summary"] = "; ".join(parts) or "текст прочитан, ключевые поля не найдены"
    return result


def parse_info_message_text(text: str) -> dict[str, Any]:
    result = {
        "doc_type": "info_message",
        "auction_date": "не указано",
        "organizer": "не указано",
        "starting_price": 0,
        "parsed_ok": False,
        "summary": "",
    }
    if not text or len(text) < 40:
        return result
    norm = re.sub(r"[ \t]+", " ", text)

    dm = re.search(
        r"(\d{1,2}[./]\d{1,2}[./]\d{2,4})",
        norm,
    )
    if dm:
        result["auction_date"] = dm.group(1)

    om = re.search(
        r"(?:организатор|арбитражн[^\n]{0,20}управля)[:\s]+([^\n]{5,120})",
        norm, re.I,
    )
    if om:
        result["organizer"] = om.group(1).strip()[:160]

    pm = re.search(r"(\d[\d\s]{5,11})\s*(?:₽|руб)", norm, re.I)
    if pm:
        result["starting_price"] = int(re.sub(r"\s", "", pm.group(1)))

    result["parsed_ok"] = any(
        x != "не указано" or result["starting_price"]
        for x in (result["auction_date"], result["organizer"])
    ) or result["starting_price"] > 0

    parts = []
    if result["auction_date"] != "не указано":
        parts.append(f"дата: {result['auction_date']}")
    if result["organizer"] != "не указано":
        parts.append(f"организатор: {result['organizer'][:60]}")
    if result["starting_price"]:
        parts.append(f"цена {result['starting_price']/1e6:.2f} млн")
    result["summary"] = "; ".join(parts) or "текст прочитан, ключевые поля не найдены"
    return result


def parse_document_content(doc_type: str, text: str, title: str = "") -> dict[str, Any]:
    if doc_type == "egrn":
        from analyzer import parse_egrn_pdf
        parsed = parse_egrn_pdf(text)
        parsed["doc_type"] = "egrn"
        return parsed
    if doc_type == "appraisal":
        from appraisal_pdf import parse_appraisal_pdf
        parsed = parse_appraisal_pdf(text)
        parsed["doc_type"] = "appraisal"
        return parsed
    if doc_type == "contract":
        return parse_contract_text(text)
    if doc_type == "application":
        return parse_application_text(text)
    if doc_type == "info_message":
        return parse_info_message_text(text)
    return {
        "doc_type": doc_type,
        "parsed_ok": bool(text and len(text) >= 80),
        "summary": (text[:200] + "…") if text and len(text) > 200 else (text or "не указано"),
    }


def format_document_extracted(entry: dict) -> str:
    if entry.get("type") == "photo":
        return "есть фото"
    extracted = entry.get("extracted") or {}
    if extracted.get("summary"):
        return extracted["summary"]
    if entry.get("text_len", 0) >= 80:
        return "текст прочитан, структура не распознана"
    if entry.get("download_ok"):
        return "скачан, текст не извлечён"
    return "не получен"


def format_lot_documents_report(lot: dict) -> str:
    lines = []
    for d in lot.get("lot_documents") or []:
        title = d.get("title") or d.get("url", "")[:60]
        dtype = d.get("type", "other")
        summary = format_document_extracted(d)
        lines.append(f"• {title} [{dtype}]: {summary}")
    if lot.get("has_photos"):
        n = sum(1 for d in lot.get("lot_documents") or [] if d.get("type") == "photo")
        if not n:
            lines.append("• фото: есть")
    return "\n".join(lines)
